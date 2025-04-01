import PyPDF2
from docx import Document
import tempfile
import os
from io import BytesIO
import pytesseract
from PIL import Image
import pdf2image

def process_document(file) -> str:
    """Process uploaded file and return text content with OCR fallback"""
    try:
        if file.name.endswith('.txt'):
            return file.read().decode('utf-8')
        
        elif file.name.endswith('.docx'):
            doc = Document(file)
            return '\n'.join([para.text for para in doc.paragraphs])
        
        elif file.name.endswith('.pdf'):
            with tempfile.NamedTemporaryFile(delete=False, suffix='.pdf') as tmp:
                tmp.write(file.getvalue())
                tmp_path = tmp.name
            
            try:
                text = ""
                with open(tmp_path, 'rb') as f:
                    reader = PyPDF2.PdfReader(f)
                    for page in reader.pages:
                        page_text = page.extract_text()
                        if page_text:
                            text += page_text + "\n"
                        else:
                            # OCR fallback for scanned PDFs
                            images = pdf2image.convert_from_path(tmp_path)
                            for image in images:
                                text += pytesseract.image_to_string(image) + "\n"
                return text
            finally:
                os.unlink(tmp_path)
        
        else:
            raise ValueError(f"Unsupported file format: {file.name}")
    
    except Exception as e:
        raise ValueError(f"Error processing document: {str(e)}")

def create_word_doc(content: str) -> BytesIO:
    """Generate Word document from text with formatting"""
    doc = Document()
    
    # Add title
    doc.add_heading('Enhanced Document', level=1)
    
    # Preserve paragraphs
    for paragraph in content.split('\n'):
        if paragraph.strip():
            doc.add_paragraph(paragraph)
    
    output = BytesIO()
    doc.save(output)
    output.seek(0)
    return output